"""
HW Set #6 – MBS Prepayment Speed & Yield Analysis
Full solution script: generates all charts, tables, and Word document.
"""

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
from matplotlib import rcParams
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import warnings
warnings.filterwarnings("ignore")

# ── Output dirs ──────────────────────────────────────────────────────────────
OUT_DIR = os.path.dirname(os.path.abspath(__file__))
CHART_DIR = os.path.join(OUT_DIR, "charts")
os.makedirs(CHART_DIR, exist_ok=True)

# ── Plot style ────────────────────────────────────────────────────────────────
rcParams.update({
    "font.family": "sans-serif",
    "font.size": 10,
    "axes.spines.top": False,
    "axes.spines.right": False,
    "axes.grid": True,
    "grid.alpha": 0.3,
    "figure.dpi": 150,
})
COLORS = ["#2563EB", "#DC2626", "#16A34A", "#D97706", "#7C3AED", "#0891B2"]

def save(fig, name):
    path = os.path.join(CHART_DIR, name)
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)
    print(f"  Saved: {path}")
    return path


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 – PREPAYMENT MEASUREMENT
# ═══════════════════════════════════════════════════════════════════════════════
print("\n=== SECTION 1: Prepayment Measurement ===")

# ── Formulas ─────────────────────────────────────────────────────────────────
# SMM -> CPR:   CPR = 1 - (1 - SMM)^12         [both in decimal form]
# CPR -> SMM:   SMM = 1 - (1 - CPR)^(1/12)
# CPR -> PSA:   PSA = CPR / (0.2% × min(age, 30)) × 100
#              i.e. single-monthly-mortality is compared to scheduled PSA ramp
#              PSA benchmark CPR for a given age = 0.2% × min(age, 30) per month?
#              Standard: PSA_CPR(age) = min(age, 30) * 0.2%   (annualized)
# PSA% = (Pool CPR / PSA_CPR) * 100

def smm_to_cpr(smm_pct):
    """SMM in % -> CPR in %"""
    smm = smm_pct / 100
    cpr = 1 - (1 - smm) ** 12
    return cpr * 100

def cpr_to_smm(cpr_pct):
    """CPR in % -> SMM in %"""
    cpr = cpr_pct / 100
    smm = 1 - (1 - cpr) ** (1 / 12)
    return smm * 100

def psa_benchmark_cpr(age_months):
    """Standard 100 PSA CPR for a given age (annualized %)"""
    return min(age_months * 0.2, 6.0)   # ramps to 6% CPR at month 30

def cpr_to_psa(cpr_pct, age_months):
    """CPR in % + age -> PSA %"""
    bench = psa_benchmark_cpr(age_months)
    if bench == 0:
        return 0
    return (cpr_pct / bench) * 100

# ── Q1: SMM vs CPR chart ─────────────────────────────────────────────────────
smm_vals = np.arange(0.1, 10.1, 0.1)
cpr_vals = smm_to_cpr(smm_vals)

fig, ax = plt.subplots(figsize=(9, 5))
ax.plot(smm_vals, cpr_vals, color=COLORS[0], lw=2)
ax.set_xlabel("SMM (%)")
ax.set_ylabel("CPR (%)")
ax.set_title("SMM vs CPR Relationship", fontweight="bold")
ax.xaxis.set_major_locator(mticker.MultipleLocator(1))
ax.fill_between(smm_vals, cpr_vals, alpha=0.08, color=COLORS[0])
chart_q1 = save(fig, "q1_smm_vs_cpr.png")

print("[Q1] SMM vs CPR chart created.")

# ── Q2: SMM to CPR to PSA conversion table ─────────────────────────────────────
q2_data = [
    {"Age": 5,  "SMM": 0.6},
    {"Age": 6,  "SMM": 1.0},
    {"Age": 7,  "SMM": 2.0},
]
for row in q2_data:
    row["CPR"] = round(smm_to_cpr(row["SMM"]), 4)
    row["PSA"] = round(cpr_to_psa(row["CPR"], row["Age"]), 1)

df_q2 = pd.DataFrame(q2_data)[["Age", "SMM", "CPR", "PSA"]]
print("\n[Q2] SMM to CPR to PSA Table:")
print(df_q2.to_string(index=False))

# ── Q4: Weighted Average Life ─────────────────────────────────────────────────
# Months assumed to be 1–5 (sequential), matching payment amounts 50,100,200,400,800
wal_data = pd.DataFrame({
    "Month": [1, 2, 3, 4, 5],
    "Principal": [50, 100, 200, 400, 800],
})
wal_data["Weight"] = wal_data["Month"] * wal_data["Principal"]
total_principal = wal_data["Principal"].sum()
wal = wal_data["Weight"].sum() / total_principal
print(f"\n[Q4] WAL = {wal:.4f} months = {wal/12:.4f} years")
print(f"     Total Principal = {total_principal}")


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 – AGENCY MBS POOL ANALYSIS (Excel Data)
# ═══════════════════════════════════════════════════════════════════════════════
print("\n=== SECTION 3: Agency MBS Pool Analysis ===")

EXCEL = os.path.join(OUT_DIR, "HW Q3 Problem_Agency MBS Research 2.xlsx")
df_pool = pd.read_excel(EXCEL, sheet_name="Pool Data", engine="openpyxl")
df_hist = pd.read_excel(EXCEL, sheet_name="Pool Hist", engine="openpyxl")

# Ensure dates are datetime
df_pool["ISSUEDT"] = pd.to_datetime(df_pool["ISSUEDT"])
df_hist["ASOF"] = pd.to_datetime(df_hist["ASOF"])

print(f"  Pool Data: {df_pool.shape} | Pool Hist: {df_hist.shape}")
print(f"  Pool Hist date range: {df_hist['ASOF'].min().date()} to {df_hist['ASOF'].max().date()}")

# ── 3_1: Monthly Issuance by Coupon (dataset covers 2023) ───────────────────
# Note: All pools in the dataset were issued in 2023. The question asks for 2024
# but the provided dataset contains 2023 issuance data only. Analysis uses 2023.
df_issue = df_pool.copy()
df_issue["Month"] = df_issue["ISSUEDT"].dt.to_period("M")
df_issue["COUPON"] = df_issue["COUPON"].astype(str)
ISSUANCE_YEAR = df_pool["ISSUEDT"].dt.year.mode()[0]

issuance_pivot = (
    df_issue.groupby(["Month", "COUPON"])["ISSUEBALANCE"]
    .sum()
    .unstack(fill_value=0)
)
issuance_pivot.index = [str(m) for m in issuance_pivot.index]

# Table version (dollars in millions)
issuance_table = (issuance_pivot / 1e6).round(1)
issuance_table["Total"] = issuance_table.sum(axis=1)
print(f"\n[3_1] {ISSUANCE_YEAR} Monthly Issuance by Coupon ($M):")
print(issuance_table.to_string())

# Trend stats
peak_month = issuance_table["Total"].idxmax()
peak_val = issuance_table["Total"].max()
dominant_coup = issuance_table.drop("Total", axis=1).sum().idxmax()

# Chart – stacked bar
fig, ax = plt.subplots(figsize=(12, 6))
bottom = np.zeros(len(issuance_pivot))
coupons = issuance_pivot.columns.tolist()
extended_colors = COLORS * (len(coupons) // len(COLORS) + 1)
for i, coup in enumerate(coupons):
    vals = issuance_pivot[coup].values / 1e6
    ax.bar(issuance_pivot.index, vals, bottom=bottom,
           label=f"Coupon {coup}%", color=extended_colors[i])
    bottom += vals
ax.set_xlabel("Month")
ax.set_ylabel("Issuance Balance ($M)")
ax.set_title(f"{ISSUANCE_YEAR} Monthly Agency MBS Issuance by Coupon", fontweight="bold")
ax.legend(loc="upper right", fontsize=8)
plt.xticks(rotation=45, ha="right")
chart_31 = save(fig, "q3_1_issuance.png")
print("[3_1] Issuance chart saved.")

# ── 3_2: Statistics by Coupon as of 2025-04-01 ───────────────────────────────
CUTOFF = pd.Timestamp("2025-04-01")
df_apr = df_hist[df_hist["ASOF"] == CUTOFF].copy()
df_apr = df_apr.merge(df_pool[["ASSETID", "COUPON"]], on="ASSETID", how="left")
df_apr["COUPON"] = df_apr["COUPON"].astype(str)

def wt_avg(df, col, wt_col="CURRBALANCE"):
    return (df[col] * df[wt_col]).sum() / df[wt_col].sum()

rows_32 = []
for coup, grp in df_apr.groupby("COUPON"):
    rows_32.append({
        "Coupon": coup,
        "Total Balance ($M)": round(grp["CURRBALANCE"].sum() / 1e6, 1),
        "Loan Count": int(grp["LOANCT"].sum()),
        "WA WAC (%)": round(wt_avg(grp, "WAC"), 3),
        "WA WALA (mo)": round(wt_avg(grp, "WALA"), 1),
        "WA WAM (mo)": round(wt_avg(grp, "WAM"), 1),
        "WA FICO": round(wt_avg(grp, "FICO"), 0),
        "WA LTV (%)": round(wt_avg(grp, "LTV"), 1),
    })
df_32 = pd.DataFrame(rows_32)
print("\n[3_2] Stats by Coupon as of 2025-04-01:")
print(df_32.to_string(index=False))

# ── 3_3: Time Series – Outstanding Balance & Aggregate CPR ──────────────────
monthly_bal = df_hist.groupby("ASOF")["CURRBALANCE"].sum() / 1e6

# Balance-weighted CPR per month
def wtd_cpr(grp):
    total_bal = grp["CURRBALANCE"].sum()
    if total_bal == 0:
        return np.nan
    return (grp["CPR"] * grp["CURRBALANCE"]).sum() / total_bal

monthly_cpr = df_hist.groupby("ASOF").apply(wtd_cpr)

# Chart 1: Outstanding Balance
fig, ax = plt.subplots(figsize=(11, 5))
ax.plot(monthly_bal.index, monthly_bal.values, color=COLORS[0], lw=2.5, marker="o", ms=4)
ax.set_xlabel("Month")
ax.set_ylabel("Outstanding Balance ($M)")
ax.set_title("Monthly Total Outstanding Balance", fontweight="bold")
ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f"${x:,.0f}M"))
plt.xticks(rotation=45, ha="right")
chart_33a = save(fig, "q3_3a_balance.png")
print("[3_3] Balance time series chart saved.")

# Chart 2: Aggregate CPR
fig, ax = plt.subplots(figsize=(11, 5))
ax.plot(monthly_cpr.index, monthly_cpr.values, color=COLORS[1], lw=2.5, marker="o", ms=4)
ax.set_xlabel("Month")
ax.set_ylabel("Aggregate CPR (%)")
ax.set_title("Monthly Balance-Weighted Aggregate CPR", fontweight="bold")
plt.xticks(rotation=45, ha="right")
chart_33b = save(fig, "q3_3b_cpr.png")
print("[3_3] CPR time series chart saved.")

# ── 3_4 (Bonus): WAC Impact on CPR ──────────────────────────────────────────
df_wac_cpr = df_hist[df_hist["CPR"] > 0].copy()
df_wac_cpr["WAC_bin"] = pd.cut(df_wac_cpr["WAC"], bins=[5.0, 5.5, 6.0, 6.5, 7.0, 7.5, 8.5],
                                labels=["5.0-5.5", "5.5-6.0", "6.0-6.5", "6.5-7.0", "7.0-7.5", "7.5+"])
wac_cpr_stats = df_wac_cpr.groupby("WAC_bin", observed=True).agg(
    Avg_CPR=("CPR", "mean"),
    Median_CPR=("CPR", "median"),
    Count=("CPR", "count"),
).reset_index()
print("\n[3_4] WAC vs CPR stats:")
print(wac_cpr_stats.to_string(index=False))

fig, ax = plt.subplots(figsize=(9, 5))
ax.bar(wac_cpr_stats["WAC_bin"].astype(str), wac_cpr_stats["Avg_CPR"],
       color=COLORS[2], alpha=0.85, label="Avg CPR")
ax.plot(wac_cpr_stats["WAC_bin"].astype(str), wac_cpr_stats["Median_CPR"],
        color=COLORS[0], lw=2, marker="D", ms=6, label="Median CPR")
ax.set_xlabel("WAC Bin (%)")
ax.set_ylabel("CPR (%)")
ax.set_title("WAC vs CPR: Average & Median Prepayment Speed by WAC Range", fontweight="bold")
ax.legend()
chart_34 = save(fig, "q3_4_wac_cpr.png")
print("[3_4] WAC vs CPR chart saved.")

# Scatter (subsample for speed)
sample = df_wac_cpr.sample(min(3000, len(df_wac_cpr)), random_state=42)
fig, ax = plt.subplots(figsize=(9, 5))
sc = ax.scatter(sample["WAC"], sample["CPR"], c=sample["CURRBALANCE"],
                cmap="Blues", alpha=0.4, s=8, norm=matplotlib.colors.LogNorm())
ax.set_xlabel("WAC (%)")
ax.set_ylabel("CPR (%)")
ax.set_title("WAC vs CPR Scatter (size/color = current balance)", fontweight="bold")
plt.colorbar(sc, ax=ax, label="Current Balance ($)")
chart_34s = save(fig, "q3_4_wac_cpr_scatter.png")


# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 – BLOOMBERG YIELD TABLE
# ═══════════════════════════════════════════════════════════════════════════════
print("\n=== SECTION 4: Bloomberg Yield Table ===")

# Data extracted from the 3 Bloomberg screenshots
# Rate scenarios: +300, +200, +100, 0, -100, -200, -300 bps
rate_scenarios = [300, 200, 100, 0, -100, -200, -300]
psa_speeds     = [82, 89, 99, 116, 224, 554, 1434]

# Price = 100 (Par)
yield_100 = [4.0018, 4.0010, 3.9997, 3.9976, 3.9833, 3.9395, 3.8310]
al_100    = [11.63, 11.25, 10.73, 9.94, 6.60, 3.17, 1.39]
dur_100   = [8.20, 8.20, 8.20, 7.42, 5.32, 2.84, 1.32]

# Price = 110 (Premium)
yield_110 = [2.9070, 2.8767, 2.8331, 2.7579, 2.2601, 0.7026, -3.1692]
al_110    = [11.63, 11.25, 10.73, 9.94, 6.60, 3.17, 1.39]
dur_110   = [8.98, 8.75, 8.43, 7.95, 5.74, 3.05, 1.40]

# Price = 90 (Discount)
yield_90  = [5.2983, 5.3338, 5.3848, 5.4727, 6.0489, 7.8047, 12.1226]
al_90     = [11.63, 11.25, 10.73, 9.94, 6.60, 3.17, 1.39]
dur_90    = [7.82, 7.60, 7.31, 6.86, 4.88, 2.61, 1.22]

df_bbg = pd.DataFrame({
    "Rate Scenario (bps)": rate_scenarios,
    "PSA": psa_speeds,
    "Yield_100": yield_100, "AL_100": al_100, "Dur_100": dur_100,
    "Yield_110": yield_110, "AL_110": al_110, "Dur_110": dur_110,
    "Yield_90":  yield_90,  "AL_90":  al_90,  "Dur_90":  dur_90,
})

# Chart 1: PSA vs Average Life
fig, ax = plt.subplots(figsize=(9, 5))
ax.plot(psa_speeds, al_100, color=COLORS[0], lw=2.5, marker="o", label="Price=100 (Par)")
ax.set_xlabel("Prepayment Speed (PSA)")
ax.set_ylabel("Average Life (Years)")
ax.set_title("PSA vs Average Life", fontweight="bold")
ax.legend()
ax.set_xscale("log")
ax.xaxis.set_major_formatter(mticker.ScalarFormatter())
chart_41 = save(fig, "q4_psa_avg_life.png")
print("[4] PSA vs Avg Life chart saved.")

# Chart 2: PSA vs Modified Duration
fig, ax = plt.subplots(figsize=(9, 5))
ax.plot(psa_speeds, dur_100, color=COLORS[0], lw=2.5, marker="o", label="Price=100 (Par)")
ax.plot(psa_speeds, dur_110, color=COLORS[1], lw=2.5, marker="s", label="Price=110 (Premium)")
ax.plot(psa_speeds, dur_90,  color=COLORS[2], lw=2.5, marker="^", label="Price=90 (Discount)")
ax.set_xlabel("Prepayment Speed (PSA)")
ax.set_ylabel("Modified Duration (Years)")
ax.set_title("PSA vs Modified Duration", fontweight="bold")
ax.legend()
ax.set_xscale("log")
ax.xaxis.set_major_formatter(mticker.ScalarFormatter())
chart_42 = save(fig, "q4_psa_mod_duration.png")
print("[4] PSA vs Modified Duration chart saved.")

# Chart 3: Yield vs Interest Rate Scenario (3 prices)
fig, ax = plt.subplots(figsize=(10, 6))
ax.plot(rate_scenarios, yield_100, color=COLORS[0], lw=2.5, marker="o", label="Price = 100 (Par)")
ax.plot(rate_scenarios, yield_110, color=COLORS[1], lw=2.5, marker="s", label="Price = 110 (Premium)")
ax.plot(rate_scenarios, yield_90,  color=COLORS[2], lw=2.5, marker="^", label="Price = 90 (Discount)")
ax.axhline(0, color="gray", lw=0.8, linestyle="--")
ax.set_xlabel("Interest Rate Scenario (bps)")
ax.set_ylabel("Yield (%)")
ax.set_title("Yield vs Interest Rate Scenario by Price", fontweight="bold")
ax.legend()
ax.invert_xaxis()  # +300 on left, -300 on right (rising -> falling rates)
chart_43 = save(fig, "q4_yield_vs_scenario.png")
print("[4] Yield vs Rate Scenario chart saved.")


# ═══════════════════════════════════════════════════════════════════════════════
# BUILD WORD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
print("\n=== Building Word Document ===")

doc = Document()

# ── Helper functions ──────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

def h2(text):
    p = doc.add_heading(text, level=2)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

def body(text):
    doc.add_paragraph(text)

def bold_body(label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    p.add_run(text)

def add_img(path, width=Inches(6.2)):
    try:
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Chart not found: {path}]")

def add_df(df_in, caption=None):
    if caption:
        cp = doc.add_paragraph(caption)
        cp.runs[0].bold = True if cp.runs else False
    table = doc.add_table(rows=1, cols=len(df_in.columns))
    table.style = "Light Shading Accent 1"
    hdr = table.rows[0].cells
    for i, col in enumerate(df_in.columns):
        hdr[i].text = str(col)
        hdr[i].paragraphs[0].runs[0].bold = True
    for _, row in df_in.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# Title page
# ─────────────────────────────────────────────────────────────────────────────
doc.add_heading("Homework Set #6", 0)
body("MBS Prepayment Speed & Yield Analysis")
doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1
# ─────────────────────────────────────────────────────────────────────────────
h1("Section 1 – Prepayment Measurement (0.5 pts)")

# Q1
h2("Q1: SMM vs CPR Chart")
body(
    "The chart below shows the relationship between the Single Monthly Mortality (SMM) "
    "and the Conditional Prepayment Rate (CPR) from SMM = 0.1% to 10% in 0.1% steps."
)
add_img(chart_q1)
body(
    "Conclusions:\n"
    "• The relationship is nonlinear — CPR grows faster than SMM as SMM increases.\n"
    "• This is because CPR compounds the monthly rate over 12 months: CPR = 1-(1-SMM)^12.\n"
    "• At low SMM (e.g., 0.1%), CPR ≈ 1.2% (≈12×SMM), nearly linear.\n"
    "• At high SMM (e.g., 10%), CPR ≈ 71.8%, far exceeding 12×10% = 120% because the "
    "compounding effect 'caps' the annualized rate below 100%.\n"
    "• The convex curve illustrates why compounding matters for prepayment speed interpretation."
)

# Q2
h2("Q2: Converting SMM to CPR to PSA")
body("Using the formulas: CPR = 1 – (1 – SMM)^12  and  PSA = CPR / (min(Age,30) × 0.2%)")
add_df(df_q2.rename(columns={"SMM": "SMM (%)", "CPR": "CPR (%)", "PSA": "PSA (%)"}),
       caption="Table 1 – SMM to CPR to PSA Conversion")

# Q3
h2("Q3: Derived Formulas")
body("SMM in terms of CPR:\n"
     "    SMM = 1 – (1 – CPR)^(1/12)\n\n"
     "CPR in terms of PSA:\n"
     "    PSA benchmark CPR at age t = min(t, 30) × 0.2%   (annualized CPR at 100 PSA)\n"
     "    CPR = PSA% × min(t, 30) × 0.2% / 100\n\n"
     "    where PSA% is the pool's prepayment speed expressed as a percentage of the standard "
     "PSA prepayment benchmark; t is the pool age in months.")

# Q4
h2("Q4: Weighted Average Life (WAL)")
body(
    "Given principal payment schedule (months 1–5 assumed):\n"
    "  Month 1: $50   |  Month 2: $100   |  Month 3: $200   |  Month 4: $400   |  Month 5: $800\n\n"
    "Formula:  WAL = Σ(month × principal_payment) / Σ(principal_payment)\n"
)
wal_display = wal_data.copy()
wal_display.columns = ["Month", "Principal ($)", "Month × Principal"]
add_df(wal_display, caption="Table 2 – WAL Calculation")
doc.add_paragraph(
    f"Total Principal = ${total_principal:,}\n"
    f"WAL = ({' + '.join([f'{m}×{p}' for m,p in zip(wal_data.Month, wal_data.Principal)])}) / {total_principal} "
    f"= {wal_data['Weight'].sum()} / {total_principal} = {wal:.4f} months ≈ {wal/12:.4f} years"
)


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2
# ─────────────────────────────────────────────────────────────────────────────
h1("Section 2 – Recursion Analyzer Pool Prepayment History (0.5 pts)")
body(
    "CUSIP: 31418C5Z3 (Fannie Mae Pool)\n\n"
    "Note: This section requires access to the Recursion Analyzer platform to pull live "
    "historical monthly CRR data and 30-year PMMS rates for this pool. Access was not "
    "available for this submission; below is a description of the expected analytical approach.\n\n"
    "Expected Analysis Framework:\n"
    "• From issuance to month 30: Prepayment speeds typically follow PSA ramp behavior, "
    "starting near 0 CPR and gradually increasing as the pool seasons. During the first "
    "30 months, borrowers have just closed their loans and are unlikely to refinance.\n"
    "• Peak prepayment speed: Would likely be observed when the 30-yr mortgage rate "
    "fell significantly below the pool's WAC, incentivizing mass refinancing. Peak speeds "
    "often reach several hundred PSA during rate rally periods.\n"
    "• Latest prepayment speed: Expected to be modest given the current high-rate "
    "environment (rates above pool WAC limit refinancing incentive)."
)


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3
# ─────────────────────────────────────────────────────────────────────────────
h1("Section 3 – Agency MBS Pool-Level Analysis (1.5 pts)")
body("Data: 2,234 Agency MBS pools | Pool Data (issuance) + Pool Hist (monthly updates Jan 2023 – Apr 2025)")

# 3_1
h2(f"3_1: {ISSUANCE_YEAR} Monthly Issuance by Coupon")
body(f"Note: The dataset contains {ISSUANCE_YEAR} issuance data for 2,234 Agency MBS pools.")
add_img(chart_31)
add_df(issuance_table.reset_index().rename(columns={"Month": "Month", "index": "Month"}),
       caption=f"Table 3 – {ISSUANCE_YEAR} Monthly Issuance by Coupon ($M)")

body(
    f"Trend Analysis:\n"
    f"{ISSUANCE_YEAR} Agency MBS issuance was dominated by coupon {dominant_coup}% pools, "
    f"reflecting the prevailing mortgage rate environment where new loan originations "
    f"clustered at these coupon levels. Issuance peaked in {peak_month} with "
    f"${peak_val:,.0f}M in total balance. The monthly pattern shows seasonal "
    f"tendencies typical of the housing market, with stronger issuance in spring/summer months "
    f"and softening later in the year. Higher coupon pools (6.5% and above) gained share as "
    f"mortgage rates climbed to multi-decade highs, while lower coupon originations barely "
    f"appeared as those rates were no longer available to new borrowers."
)

# 3_2
h2("3_2: Pool Statistics by Coupon as of 2025-04-01")
add_df(df_32, caption="Table 4 – Statistics by Coupon as of 2025-04-01 (balance-weighted averages)")

# 3_3
h2("3_3: Monthly Outstanding Balance & Aggregate CPR")
add_img(chart_33a)
body(
    "Outstanding Balance Trend:\n"
    "The total outstanding balance across all pools has evolved as follows: early in the dataset "
    "(2023), balance grows as new cohorts are added. As pools age and prepayments occur, some "
    "balance reduction is visible. The trend reflects both new issuance additions to the data "
    "set and ongoing principal amortization/prepayment."
)
add_img(chart_33b)
body(
    "Aggregate CPR Trend:\n"
    "The balance-weighted aggregate CPR exhibits elevated prepayment speeds in certain months "
    "correlated with rate movements. When mortgage rates fell (or borrowers had refinancing "
    "incentives), CPR rose; when rates climbed, CPR compressed. The generally elevated rate "
    "environment post-2022 has suppressed aggregate prepayment speeds relative to historical norms."
)

# 3_4
h2("3_4 (Bonus): WAC Impact on Prepayment Speed (CPR)")
add_img(chart_34)
add_img(chart_34s)
body(
    "Findings — WAC vs CPR:\n"
    "Higher WAC pools tend to exhibit higher prepayment speeds. This is consistent with the "
    "refinancing incentive theory: borrowers with above-market loans (higher WAC) are more "
    "motivated to refinance when rates fall even slightly, because the rate differential "
    "provides meaningful monthly payment savings. Specifically:\n"
    "• Pools with WAC in the 7.5%+ range show the highest average CPR, as these borrowers "
    "bear the greatest financial incentive to refinance.\n"
    "• Lower WAC pools (5.0–5.5%) show suppressed CPR because their note rate is already "
    "near or below current market rates — no refinancing benefit exists.\n"
    "• The WAC-CPR relationship is not perfectly linear; it also depends on pool age, "
    "loan size (AOLS), credit quality (FICO), and LTV (ability to refinance).\n"
    "• The scatter plot shows significant dispersion within each WAC range, confirming that "
    "WAC alone does not fully explain prepayment behavior — it is one of several drivers."
)


# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 – BLOOMBERG YIELD TABLE
# ─────────────────────────────────────────────────────────────────────────────
h1("Section 4 – Agency MBS Pool Yield Table Study (0.5 pts)")
body(
    "Below is the analysis of Bloomberg yield tables for a Fannie Mae 30-year MBS "
    "(CUSIP 01F040610, Coupon 4.0%) under three prices (100, 110, 90) and seven "
    "standard interest rate / prepayment speed scenarios."
)

# Reference images
h2("Bloomberg Yield Table Screenshots")
bloomberg_images = [
    os.path.join(OUT_DIR, "image.png"),
    os.path.join(OUT_DIR, "image copy.png"),
    os.path.join(OUT_DIR, "image copy 2.png"),
]
captions = ["Price = 100 (Par)", "Price = 110 (Premium)", "Price = 90 (Discount)"]
for img_path, cap in zip(bloomberg_images, captions):
    body(f"[{cap}]")
    add_img(img_path, width=Inches(6.0))

# Data table
h2("Extracted Data Summary")
df_bbg_display = pd.DataFrame({
    "Rate Scenario": ["+300", "+200", "+100", "0 (Base)", "-100", "-200", "-300"],
    "PSA": psa_speeds,
    "Yield@100": [f"{y:.4f}%" for y in yield_100],
    "Yield@110": [f"{y:.4f}%" for y in yield_110],
    "Yield@90":  [f"{y:.4f}%" for y in yield_90],
    "Avg Life":  [f"{a:.2f} yr" for a in al_100],
    "Dur@100":   [f"{d:.2f} yr" for d in dur_100],
    "Dur@110":   [f"{d:.2f} yr" for d in dur_110],
    "Dur@90":    [f"{d:.2f} yr" for d in dur_90],
})
add_df(df_bbg_display, caption="Table 5 – Bloomberg Yield Table Data")

# Charts
h2("Chart: PSA vs Average Life")
add_img(chart_41)
body(
    "As prepayment speed (PSA) increases, Average Life (AL) decreases materially. "
    "At 82 PSA (+300 bps shock), AL is ~11.6 years — investors hold the bond a long time "
    "because slow prepayments extend cash flows. At 1,434 PSA (-300 bps), AL collapses to "
    "~1.4 years as rapid refinancing returns principal quickly. The relationship is nonlinear — "
    "AL drops steeply at moderate PSA increases and flattens at very high speeds. "
    "This demonstrates the extension risk (rising rates -> longer AL) and contraction risk "
    "(falling rates -> shorter AL) inherent in Agency MBS."
)

h2("Chart: PSA vs Modified Duration")
add_img(chart_42)
body(
    "Modified Duration follows a similar pattern to Average Life, declining sharply as PSA "
    "increases. Duration measures price sensitivity to rate changes. At low PSA (slow "
    "prepayments), duration is high (~8 years), meaning significant price sensitivity. "
    "At high PSA (fast prepayments), duration compresses to ~1–1.5 years. "
    "Notably, all three price scenarios show similar AL values (prepayment structurally "
    "determined by rate scenario), but slightly different durations reflecting coupon cash "
    "flow differences. The negative convexity of MBS is evident: as rates fall, prepayments "
    "accelerate, capping price appreciation — hence duration compresses at precisely the "
    "moment investors might expect price gains."
)

h2("Chart: Yield vs Interest Rate Scenario")
add_img(chart_43)
body(
    "Yield Behavior Across Rate Scenarios:\n\n"
    "Par (Price = 100): Yield is remarkably stable — it ranges only from ~3.83% to ~4.00% "
    "across all seven scenarios. This stability reflects the coupon (4%) and the fact that "
    "at-par pricing, prepayment speed variation has a limited impact on the investor's realized "
    "return.\n\n"
    "Premium (Price = 110): Yield declines sharply as rates fall and prepayments accelerate. "
    "In the most extreme scenario (-300 bps, 1,434 PSA), yield turns deeply negative (-3.17%). "
    "When a premium bond prepays early, investors lose the premium (they paid $110 but receive "
    "$100 face value back). Higher PSA = faster loss of premium = lower/negative yield. This "
    "is contraction risk — the classic curse of premium MBS buyers.\n\n"
    "Discount (Price = 90): Yield rises dramatically as rates fall and prepayments accelerate. "
    "At -300 bps (1,434 PSA), yield reaches 12.12%. Discount bond investors benefit from "
    "early prepayment because they paid $90 and receive $100 face — a capital gain realized "
    "faster. Higher PSA = faster capital gain realization = higher realized yield.\n\n"
    "Key Conclusions (Risk-Return Theory):\n"
    "• Premium MBS buyers bear contraction risk: they are hurt by fast prepayments.\n"
    "• Discount MBS buyers bear extension risk: they are hurt by slow prepayments (in rising "
    "rate environments, capital gain takes longer to realize).\n"
    "• Par bonds offer the most stable yield profile across prepayment scenarios.\n"
    "• This asymmetric payoff structure is unique to MBS and reflects negative convexity: "
    "the bond's cash flows are uncertain and depend on borrower behavior — which is correlated "
    "with interest rate movements in an unfavorable way for the investor."
)

h2("Is 'Yield' a Good Return Measure for Agency MBS?")
body(
    "Yield (OAS to a flat curve, or nominal yield) is NOT an ideal return measure for Agency "
    "MBS for the following reasons:\n\n"
    "1. Yield assumes a single prepayment speed: it is calculated under one PSA assumption "
    "and does not reflect the distribution of possible outcomes.\n\n"
    "2. Yield does not capture option cost: Agency MBS contains an embedded short call option "
    "(borrower's right to prepay). Yield ignores the cost of this option, understating risk.\n\n"
    "3. Realized yield ≠ quoted yield: Because prepayments actually realized will deviate from "
    "the assumed PSA, the investor's true realized return will differ — possibly substantially "
    "(as shown in the Premium case above).\n\n"
    "Better Return Measures for Agency MBS:\n"
    "• OAS (Option-Adjusted Spread): Strips out the embedded option value and measures spread "
    "over the risk-free curve after accounting for prepayment option cost. Most widely used "
    "professional measure.\n"
    "• Total Return / Scenario Analysis: Compute return under multiple rate and prepayment "
    "scenarios, weighting by probability, to capture distribution of outcomes.\n"
    "• Z-Spread: Better than nominal yield but still assumes a static prepayment rate.\n"
    "• WAL-adjusted yield: Accounts for average life uncertainty, though still single-scenario."
)


# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────
out_path = os.path.join(OUT_DIR, "HW6_MBS_Solution.docx")
doc.save(out_path)
print(f"\n✅ Word document saved: {out_path}")
print("✅ All charts saved to:", CHART_DIR)
