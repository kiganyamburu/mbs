"""
Update Section 2 in the existing Word document HW6_MBS_Solution.docx
with real CPR data derived from PoolTalk factors.
"""
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUT_DIR = r"c:\Users\nduta\OneDrive\Desktop\New folder (2)"
CHART_DIR = os.path.join(OUT_DIR, "charts")

doc = Document(os.path.join(OUT_DIR, "HW6_MBS_Solution.docx"))

# Find the Section 2 heading paragraph index
sec2_idx = None
for i, para in enumerate(doc.paragraphs):
    if "Section 2" in para.text and "Recursion" in para.text:
        sec2_idx = i
        break

if sec2_idx is None:
    print("ERROR: Could not find Section 2 heading")
    exit(1)

print(f"Section 2 heading at paragraph index {sec2_idx}: '{doc.paragraphs[sec2_idx].text}'")

# Find end of Section 2 (start of Section 3 heading)
sec3_idx = None
for i in range(sec2_idx+1, len(doc.paragraphs)):
    if "Section 3" in doc.paragraphs[i].text:
        sec3_idx = i
        break

print(f"Section 3 heading at paragraph index {sec3_idx}")

# Delete all Section 2 body paragraphs (between sec2_idx and sec3_idx)
# We'll replace them with updated content. Work backwards to keep indices valid.
body = doc.element.body
paras = list(body)  # all block-level elements

# Get the XML elements for the range to remove
# We need to identify all child elements between the two headings
all_body_children = list(body)
heading2_el = doc.paragraphs[sec2_idx]._element
heading3_el = doc.paragraphs[sec3_idx]._element

# Collect elements to remove
to_remove = []
found = False
for el in all_body_children:
    if el is heading2_el:
        found = True
        continue
    if el is heading3_el:
        break
    if found:
        to_remove.append(el)

print(f"Removing {len(to_remove)} elements from Section 2")
for el in to_remove:
    body.remove(el)

# Now insert new content before Section 3 heading
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def insert_paragraph_before(ref_para, text, style="Normal"):
    """Insert a new paragraph before ref_para"""
    new_para = OxmlElement("w:p")
    ref_para._element.addprevious(new_para)
    doc_para = doc.paragraphs[doc.paragraphs.index(ref_para) - 1]
    # Actually, just add runs to the new element
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_para, doc)
    p.style = doc.styles[style]
    p.text = text
    return p

def add_before_sec3(text, style="Normal"):
    """Add paragraph before the Section 3 heading"""
    heading3 = doc.paragraphs[next(i for i, p in enumerate(doc.paragraphs) if "Section 3" in p.text)]
    new_para = OxmlElement("w:p")
    heading3._element.addprevious(new_para)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_para, doc)
    p.style = doc.styles[style]
    p.text = text
    return p

def add_image_before_sec3(img_path, width=Inches(6.2)):
    """Add picture before Section 3"""
    heading3 = doc.paragraphs[next(i for i, p in enumerate(doc.paragraphs) if "Section 3" in p.text)]
    # Add a new paragraph for the image
    new_para = OxmlElement("w:p")
    heading3._element.addprevious(new_para)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_para, doc)
    run = p.add_run()
    run.add_picture(img_path, width=width)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_heading2_before_sec3(text):
    """Add heading level 2 before Section 3"""
    heading3 = doc.paragraphs[next(i for i, p in enumerate(doc.paragraphs) if "Section 3" in p.text)]
    new_para = OxmlElement("w:p")
    heading3._element.addprevious(new_para)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_para, doc)
    p.style = doc.styles["Heading 2"]
    p.text = text
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    return p

# Build Section 2 content
add_heading2_before_sec3("Pool Overview")
add_before_sec3(
    "CUSIP: 31418C5Z3  |  Pool: MA3563  |  Issuer: Fannie Mae (UMBS)\n"
    "Issue Date: December 2018  |  Coupon: 4.000%  |  Original UPB: $6.39 Billion\n"
    "Maturity: January 2049  |  Current UPB (Mar 2026): $425.9M  |  Remaining Factor: 0.0667"
)

add_heading2_before_sec3("Historical Monthly CPR vs. 30-Year Primary Mortgage Rate")
add_before_sec3(
    "Data source: Monthly security factors from Fannie Mae PoolTalk (public disclosure tool). "
    "CPR derived from sequential factor changes using: CPR = 1-(Factor_t/Factor_{t-1})^12. "
    "30-year mortgage rate (PMMS) sourced from FRED (series: MORTGAGE30US)."
)
add_image_before_sec3(os.path.join(CHART_DIR, "sec2_cusip_cpr_pmms.png"), width=Inches(6.5))

add_heading2_before_sec3("Pattern Descriptions")

add_before_sec3(
    "1) ISSUANCE (Dec 2018) TO MONTH 30 (May 2021):\n\n"
    "This pool exhibited rapid prepayment escalation from issuance, far exceeding the normal PSA ramp pattern. "
    "The pool was issued in Dec 2018 when the 30-yr mortgage rate was ~4.6%, just above the pool coupon of 4.0%, "
    "meaning borrowers had limited initial refinancing incentive. Early CPR was modest (~4-9%).\n\n"
    "However, rates began falling through 2019 as the Fed eased, and CPR accelerated sharply — reaching "
    "30-50% CPR by mid-2019 as the PMMS dropped toward and then below 3.8%. The COVID-19 shock in early 2020 "
    "sent rates to historic lows (sub-3%), triggering a massive refinancing wave. By month 19 (July 2020), "
    "CPR peaked at 66.05% — one of the fastest prepayment speeds for this pool. "
    "Through month 30, the pool prepaid at extremely high speeds consistently (40-66% CPR range), "
    "burning through over 70% of its original balance."
)

add_before_sec3(
    "2) PEAK PREPAYMENT SPEED:\n\n"
    "Peak CPR: 66.05%  |  Month: July 2020  |  Pool Age: 19 months\n"
    "30-yr PMMS at peak: ~3.13% — approximately 87 bps below the pool coupon of 4.0%.\n\n"
    "This represents the most extreme refinancing wave. With market rates nearly 90 bps below the pool's "
    "WAC, virtually all borrowers in the pool had strong economic incentive to refinance. The pool's "
    "large original balance ($6.39B) also reflects thousands of loans — statistically, the mass "
    "refinancing translated directly into a near-record CPR. The peak period (mid-2020 through mid-2021) "
    "saw sustained CPR of 50-66%, accelerating the pool's amortization dramatically."
)

add_before_sec3(
    "3) LATEST PREPAYMENT SPEED (Most Recent Available: March 2026):\n\n"
    "Latest CPR: 8.74%  |  Date: March 2026  |  Pool Age: 87 months\n"
    "30-yr PMMS (Mar 2026): ~6.0% — approximately 200 bps ABOVE the pool coupon of 4.0%.\n\n"
    "Prepayment speeds have collapsed to single digits for a very logical reason: the current mortgage "
    "rate environment (~6%) is far above the pool's coupon (4.0%). Borrowers who remain in this pool "
    "already have a 4% mortgage — refinancing would mean taking on a new loan at ~6%, resulting in "
    "significantly higher monthly payments. This creates a 'lock-in effect' where rational borrowers "
    "have no incentive to prepay voluntarily.\n\n"
    "The remaining ~8.74% CPR reflects only involuntary prepayments: defaults, home sales (job changes, "
    "divorces, deaths), and curtailments. The pool has shrunk to just 6.67% of its original balance "
    "($425.9M remaining of $6.39B original), meaning the vast majority of loans already prepaid "
    "during the 2019-2022 refinancing wave."
)

doc.save(os.path.join(OUT_DIR, "HW6_MBS_Solution.docx"))
print("Word document updated successfully with real Section 2 data!")
